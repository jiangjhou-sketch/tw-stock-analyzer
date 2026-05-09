#!/usr/bin/env python3
"""
台股均量追蹤系統 v5
資料來源：Yahoo Finance (優先) + TWSE + TPEX OpenAPI
技術指標：KD / MACD / 布林通道
yfinance 1.x 相容，修正 MultiIndex 欄位結構問題
"""

import requests, time, re, json, threading, uuid, os, io, math
from datetime import datetime
from flask import Flask, jsonify, send_from_directory, request, send_file
import yfinance as yf
import pandas as pd
import numpy as np

# ── 安全 JSON Provider（NaN/Inf → null，numpy 型別轉換）──
from flask.json.provider import DefaultJSONProvider

class _SafeJSONProvider(DefaultJSONProvider):
    @staticmethod
    def default(o):
        if isinstance(o, np.integer):  return int(o)
        if isinstance(o, np.floating): return None if (np.isnan(o) or np.isinf(o)) else float(o)
        if isinstance(o, np.bool_):    return bool(o)
        if isinstance(o, np.ndarray):  return o.tolist()
        return DefaultJSONProvider.default(o)

    def dumps(self, obj, **kw):
        kw.setdefault('allow_nan', False)
        try:
            return super().dumps(obj, **kw)
        except (ValueError, TypeError):
            raw = super().dumps(obj, allow_nan=True)
            raw = re.sub(r'\bNaN\b',       'null', raw)
            raw = re.sub(r'\bInfinity\b',  'null', raw)
            raw = re.sub(r'\b-Infinity\b', 'null', raw)
            return raw

app = Flask(__name__, static_folder='static')
app.json_provider_class = _SafeJSONProvider
app.json = _SafeJSONProvider(app)

BASE_HEADERS = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36',
    'Accept': 'application/json, text/javascript, */*; q=0.01',
    'Accept-Language': 'zh-TW,zh;q=0.9,en-US;q=0.8',
}

tasks           = {}
data_source_log = {}

@app.errorhandler(404)
def not_found(e): return jsonify({'error': '找不到此路由', 'status': 404}), 404

@app.errorhandler(500)
def server_error(e): return jsonify({'error': f'伺服器錯誤：{e}', 'status': 500}), 500

@app.after_request
def ensure_json_on_error(resp):
    if resp.status_code in (404, 500) and 'application/json' not in resp.content_type:
        resp.data = json.dumps({'error': f'HTTP {resp.status_code}'})
        resp.content_type = 'application/json'
    return resp


# ══════════════════════════════════════════════════════
# 代碼 & 名稱對照表
# ══════════════════════════════════════════════════════
def _build_code_lists():
    try:
        import twstock
        twse  = sorted(k for k, v in twstock.codes.items()
                       if k.isdigit() and len(k) == 4 and '上市' in (getattr(v,'market','') or ''))
        tpex  = sorted(k for k, v in twstock.codes.items()
                       if k.isdigit() and len(k) == 4 and getattr(v,'market','') == '上櫃')
        names = {k: v.name for k, v in twstock.codes.items() if k.isdigit() and len(k) <= 5}
        print(f'[twstock] 上市 {len(twse)} 支, 上櫃 {len(tpex)} 支, 名稱 {len(names)} 筆')
        return twse, tpex, names
    except Exception as e:
        print(f'[twstock] 失敗: {e}，使用內建清單')
        return [], [], {}

_FALLBACK_TWSE, _FALLBACK_TPEX, _TW_NAMES = _build_code_lists()

# 內建精簡清單（twstock 失敗時備用）
if not _FALLBACK_TWSE:
    _FALLBACK_TWSE = [
        '0050','0056','1101','1102','1216','1301','1303','1326','1402',
        '2002','2105','2204','2207','2301','2303','2308','2317','2323',
        '2330','2337','2344','2347','2351','2353','2354','2356','2357',
        '2360','2363','2376','2377','2379','2382','2385','2388','2395',
        '2399','2401','2402','2404','2405','2406','2408','2409','2412',
        '2414','2417','2419','2420','2423','2426','2428','2430','2432',
        '2433','2436','2448','2449','2450','2454','2458','2461','2464',
        '2465','2466','2467','2474','2475','2478','2484','2486','2492',
        '2498','2603','2604','2605','2606','2609','2610','2611','2615',
        '2618','2801','2823','2834','2836','2838','2841','2880','2881',
        '2882','2883','2884','2885','2886','2887','2888','2889','2891',
        '2892','2912','3008','3034','3035','3036','3045','3057','3094',
        '3105','3130','3231','3293','3324','3338','3374','3481','3533',
        '3605','3702','3711','3714','4904','4938','5701','5820','5876',
        '5880','6409','6415','6446','6488','6525','6547','6548','6669',
        '6770','8046','9910','9945',
    ]
if not _FALLBACK_TPEX:
    _FALLBACK_TPEX = [
        '1264','1268','1565','1569','1570','1580','1584','1586','2383',
        '3006','3042','3088','3149','3189','3209','3227','3242','3259',
        '3276','3287','3289','3305','3311','3315','3317','3321','3325',
        '3330','3413','3416','3419','3432','3443','3450','3462','3494',
        '3501','3504','3508','3513','3515','3521','3526','4106','4142',
        '4153','4168','4174','4175','4176','4537','4530','4552','4720',
        '4726','4743','4744','5234','6005','6128','6147','6150','6153',
        '6160','6161','6163','6165','6167','6170','6172','6176','6178',
        '6180','6182','6184','6186','6188','6190','6192','6194','6196',
        '6271','6472','6525','8454',
    ]


# ══════════════════════════════════════════════════════
# yfinance 1.x 相容下載（核心修正）
# ══════════════════════════════════════════════════════
def _yf_download_safe(symbols, period='90d'):
    """
    yfinance 1.x 正確下載。
    yfinance >= 0.2.38 預設 multi_level_index=True：
      columns = (Price, Ticker)，正確取法：df['Close']['SYM']
    舊版 (group_by='ticker')：df['SYM']['Close']
    兩種都自動偵測並處理。
    回傳 {symbol: DataFrame(Open/High/Low/Close/Volume)}
    """
    if isinstance(symbols, str):
        symbols = [symbols]
    symbols = list(symbols)
    if not symbols:
        return {}

    result = {}
    try:
        df = yf.download(
            symbols, period=period,
            auto_adjust=True, progress=False,
            timeout=60, threads=False,
        )
        if df is None or df.empty:
            return result

        if not isinstance(df.columns, pd.MultiIndex):
            # flat columns（可能是單股舊版）
            if 'Close' in df.columns and len(symbols) == 1:
                flat = df.dropna(subset=['Close'])
                if not flat.empty:
                    result[symbols[0]] = flat[
                        [c for c in ['Open','High','Low','Close','Volume'] if c in flat.columns]
                    ]
            return result

        lvl0 = list(df.columns.get_level_values(0).unique())
        lvl1 = list(df.columns.get_level_values(1).unique())
        price_first = 'Close' in lvl0  # True=新版(Price,Ticker)，False=舊版(Ticker,Price)

        for sym in symbols:
            try:
                if price_first:
                    if sym not in lvl1:
                        continue
                    flat = pd.DataFrame({
                        col: df[col][sym]
                        for col in ['Open','High','Low','Close','Volume']
                        if col in lvl0
                    }).dropna(subset=['Close'])
                else:
                    if sym not in lvl0:
                        continue
                    flat = pd.DataFrame({
                        col: df[sym][col]
                        for col in ['Open','High','Low','Close','Volume']
                        if col in df[sym].columns
                    }).dropna(subset=['Close'])

                if not flat.empty and len(flat) >= 5:
                    result[sym] = flat
            except Exception:
                continue

    except Exception as e:
        print(f'  [yf_download_safe] 失敗: {e}')

    return result


def get_ohlcv(symbol, days=90):
    """取得單股 OHLCV，自動嘗試 .TW / .TWO 切換"""
    for sym in [symbol,
                symbol.replace('.TW','.TWO') if symbol.endswith('.TW')
                else symbol.replace('.TWO','.TW')]:
        data = _yf_download_safe([sym], period=f'{days}d')
        if sym in data:
            return data[sym], sym
    return None, symbol


# ══════════════════════════════════════════════════════
# Yahoo Finance 漲幅排行（主要來源）
# ══════════════════════════════════════════════════════
def _get_yf_crumb():
    try:
        sess = requests.Session()
        sess.get('https://tw.stock.yahoo.com/', timeout=8,
                 headers={'User-Agent': BASE_HEADERS['User-Agent']})
        r = sess.get('https://query1.finance.yahoo.com/v1/test/getcrumb',
                     timeout=8, headers={**BASE_HEADERS, 'Referer': 'https://tw.stock.yahoo.com/'})
        if r.status_code == 200 and 0 < len(r.text) < 50:
            return sess, r.text.strip()
    except Exception:
        pass
    return requests.Session(), None


def _yahoo_tw_gainers(market='twse', top_n=100):
    """Yahoo Finance 台股漲幅排行（crumb + 多端點策略）"""
    market_tag  = '.TW' if market == 'twse' else '.TWO'
    market_name = '上市' if market == 'twse' else '上櫃'
    H = {**BASE_HEADERS, 'Referer': 'https://tw.stock.yahoo.com/',
         'Origin': 'https://tw.stock.yahoo.com'}

    def _parse(quotes):
        out = []
        for q in quotes:
            sym = str(q.get('symbol', ''))
            if not sym.endswith(market_tag): continue
            code = sym.replace(market_tag, '')
            if not re.match(r'^\d{4,5}$', code): continue
            chg = round(float(q.get('regularMarketChangePercent', 0) or 0), 2)
            if chg <= 0: continue
            out.append({
                'code': code, 'symbol': sym,
                'name': q.get('longName') or q.get('shortName') or _TW_NAMES.get(code, code),
                'change_pct': chg,
                'price':  float(q.get('regularMarketPrice', 0) or 0),
                'volume': int(q.get('regularMarketVolume', 0) or 0),
                'market': market_name,
            })
        return out

    sess, crumb = _get_yf_crumb()
    for host in ['query1', 'query2']:
        for use_post in [False, True]:
            try:
                params = {'scrIds': 'day_gainers', 'count': top_n*2, 'region': 'TW', 'lang': 'zh-TW'}
                if crumb: params['crumb'] = crumb
                if not use_post:
                    resp = sess.get(
                        f'https://{host}.finance.yahoo.com/v1/finance/screener/predefined/saved',
                        headers=H, params=params, timeout=15)
                else:
                    body = {'size': top_n*2, 'offset': 0,
                            'sortField': 'percentchange', 'sortType': 'DESC',
                            'quoteType': 'EQUITY',
                            'query': {'operator': 'AND', 'operands': [
                                {'operator': 'EQ', 'operands': ['region', 'tw']},
                                {'operator': 'GT', 'operands': ['percentchange', 0]},
                            ]},
                            'userId': '', 'userIdType': 'guid'}
                    resp = sess.post(
                        f'https://{host}.finance.yahoo.com/v1/finance/screener',
                        json=body, headers={**H, 'Content-Type': 'application/json'}, timeout=15)
                if resp.status_code == 200:
                    quotes = resp.json().get('finance',{}).get('result',[{}])[0].get('quotes',[])
                    stocks = _parse(quotes)
                    if stocks:
                        print(f'[Yahoo {host} {"POST" if use_post else "GET"}] {market_name} {len(stocks)} 支')
                        return stocks
            except Exception as e:
                print(f'[Yahoo {host}] {e}')
    return []


# ══════════════════════════════════════════════════════
# TWSE / TPEX 官方 API
# ══════════════════════════════════════════════════════
def _parse_twse_row(row):
    try:
        code = str(row[0]).strip()
        if not re.match(r'^\d{4}$', code): return None
        close = float(str(row[7]).replace(',','').strip())
        change = float(str(row[8]).replace(',','').strip())
        prev = close - change
        if prev <= 0: return None
        chg = round((change/prev)*100, 2)
        if chg <= 0: return None
        return {'code': code, 'symbol': f'{code}.TW', 'name': str(row[1]).strip(),
                'change_pct': chg, 'price': close,
                'volume': int(str(row[2]).replace(',','')), 'market': '上市'}
    except Exception: return None


def _twse_direct_api():
    H = {**BASE_HEADERS, 'Referer': 'https://www.twse.com.tw/',
         'X-Requested-With': 'XMLHttpRequest'}
    for url, params in [
        ('https://www.twse.com.tw/rwd/zh/afterTrading/STOCK_DAY_ALL', {'response':'json'}),
        ('https://www.twse.com.tw/exchangeReport/STOCK_DAY_ALL',       {'response':'json'}),
        ('https://openapi.twse.com.tw/v1/exchangeReport/STOCK_DAY_ALL', {}),
    ]:
        try:
            sess = requests.Session()
            sess.get('https://www.twse.com.tw/zh/', headers=H, timeout=8)
            resp = sess.get(url, params=params, headers=H, timeout=20)
            if resp.status_code != 200: continue
            data = resp.json()
            rows = data if isinstance(data, list) else data.get('data', [])
            if not rows: continue
            if isinstance(rows[0], dict):
                stocks = []
                for item in rows:
                    try:
                        code = str(item.get('Code','')).strip()
                        if not re.match(r'^\d{4}$', code): continue
                        close = float(str(item.get('ClosingPrice','0')).replace(',','') or 0)
                        change = float(str(item.get('Change','0')).replace(',','') or 0)
                        prev = close - change
                        if prev <= 0: continue
                        chg = round((change/prev)*100, 2)
                        if chg <= 0: continue
                        stocks.append({'code': code, 'symbol': f'{code}.TW',
                                       'name': str(item.get('Name', code)),
                                       'change_pct': chg, 'price': close,
                                       'volume': int(str(item.get('TradeVolume','0')).replace(',','') or 0),
                                       'market': '上市'})
                    except Exception: continue
            else:
                stocks = [s for row in rows for s in [_parse_twse_row(row)] if s]
            if stocks:
                print(f'[TWSE] {url[:45]} → {len(stocks)} 支')
                return stocks
        except Exception: continue
    return []


def _yf_batch_gainers(codes, market_tag, market_name, min_chg=1.0):
    """yfinance 批次取漲幅，用前收正確算法（備援）"""
    codes   = list(dict.fromkeys(c for c in codes if re.match(r'^\d{4,5}$', c)))
    symbols = [f'{c}{market_tag}' for c in codes]
    result  = {}
    for i in range(0, len(symbols), 25):
        data = _yf_download_safe(symbols[i:i+25], period='5d')
        for sym, df in data.items():
            try:
                cl = df['Close'].dropna()
                vo = df['Volume'].dropna()
                if len(cl) < 2: continue
                c_now, c_prev = float(cl.iloc[-1]), float(cl.iloc[-2])
                if c_prev <= 0: continue
                chg = round((c_now - c_prev) / c_prev * 100, 2)
                if chg < min_chg: continue
                code = sym.replace(market_tag, '')
                result[sym] = {'code': code, 'symbol': sym, 'name': _TW_NAMES.get(code, code),
                               'change_pct': chg, 'price': c_now,
                               'volume': int(vo.iloc[-1]) if not vo.empty else 0,
                               'market': market_name}
            except Exception: continue
        time.sleep(0.2)
    stocks = list(result.values())
    stocks.sort(key=lambda x: -x['change_pct'])
    print(f'[yf batch {market_name}] {len(stocks)} 支 (≥{min_chg}%)')
    return stocks


def get_twse_stocks():
    for fn in [lambda: _yahoo_tw_gainers('twse'),
               _twse_direct_api,
               lambda: _yf_batch_gainers(_FALLBACK_TWSE, '.TW', '上市')]:
        s = fn()
        if s: return s
    return []


def get_tpex_stocks():
    stocks = _yahoo_tw_gainers('tpex')
    if stocks: return stocks
    try:
        resp = requests.get('https://www.tpex.org.tw/openapi/v1/tpex_mainboard_quotes',
                            headers=BASE_HEADERS, timeout=20)
        resp.raise_for_status()
        out = []
        for item in resp.json():
            try:
                code = str(item.get('SecuritiesCompanyCode','')).strip()
                if not re.match(r'^\d{4,5}$', code): continue
                close = float(str(item.get('Close','0')).replace(',','') or 0)
                change = float(str(item.get('Change','0')).replace(',','') or 0)
                if close <= 0: continue
                prev = close - change
                if prev <= 0: continue
                chg = round((change/prev)*100, 2)
                if chg <= 0: continue
                out.append({'code': code, 'symbol': f'{code}.TWO',
                            'name': str(item.get('CompanyName','')).strip(),
                            'change_pct': chg, 'price': close,
                            'volume': int(float(str(item.get('TradeVolume','0')).replace(',','') or 0)),
                            'market': '上櫃'})
            except Exception: continue
        if out:
            print(f'[TPEX OpenAPI] {len(out)} 支')
            return out
    except Exception as e:
        print(f'[TPEX OpenAPI] 失敗: {e}')
    return _yf_batch_gainers(_FALLBACK_TPEX, '.TWO', '上櫃')


def _get_all_market_stubs(market='twse'):
    codes, tag, mkt = ((_FALLBACK_TWSE, '.TW', '上市') if market == 'twse'
                       else (_FALLBACK_TPEX, '.TWO', '上櫃'))
    stubs = [{'code': c, 'symbol': f'{c}{tag}', 'name': _TW_NAMES.get(c, c),
               'change_pct': 0.0, 'price': 0.0, 'volume': 0, 'market': mkt}
             for c in dict.fromkeys(codes)]
    print(f'[全市場] {mkt} {len(stubs)} 支')
    return stubs


def get_ranking_stocks(top_n=100, mode='both'):
    if mode == 'all_twse': return _get_all_market_stubs('twse')
    if mode == 'all_tpex': return _get_all_market_stubs('tpex')
    if mode == 'all_both': return _get_all_market_stubs('twse') + _get_all_market_stubs('tpex')
    if mode == 'twse':
        s = get_twse_stocks(); s.sort(key=lambda x:-x['change_pct']); return s[:top_n]
    if mode == 'tpex':
        s = get_tpex_stocks(); s.sort(key=lambda x:-x['change_pct']); return s[:top_n]
    if mode == 'combined':
        s = get_twse_stocks()+get_tpex_stocks(); s.sort(key=lambda x:-x['change_pct']); return s[:top_n]
    tw = get_twse_stocks(); tw.sort(key=lambda x:-x['change_pct'])
    tp = get_tpex_stocks(); tp.sort(key=lambda x:-x['change_pct'])
    return tw[:top_n] + tp[:top_n]


# ══════════════════════════════════════════════════════
# 技術指標
# ══════════════════════════════════════════════════════
def _fp(v):
    try:
        f = float(v)
        return None if (math.isnan(f) or math.isinf(f)) else f
    except Exception: return None

def _to_python(v):
    if isinstance(v, np.integer):  return int(v)
    if isinstance(v, np.floating):
        f = float(v); return None if (math.isnan(f) or math.isinf(f)) else f
    if isinstance(v, np.bool_):    return bool(v)
    if isinstance(v, float):       return None if (math.isnan(v) or math.isinf(v)) else v
    return v


def calc_kd(high, low, close, n=9, m=3):
    low_n=low.rolling(n).min(); high_n=high.rolling(n).max(); denom=high_n-low_n
    rsv=pd.Series(np.where(denom==0,50.0,(close-low_n)/denom*100),index=close.index)
    K=pd.Series(50.0,index=close.index,dtype=float)
    D=pd.Series(50.0,index=close.index,dtype=float)
    for i in range(1,len(rsv)):
        K.iloc[i]=(m-1)/m*K.iloc[i-1]+1/m*rsv.iloc[i]
        D.iloc[i]=(m-1)/m*D.iloc[i-1]+1/m*K.iloc[i]
    k_cur,k_prv=K.iloc[-1],K.iloc[-2]; d_cur,d_prv=D.iloc[-1],D.iloc[-2]
    golden=bool((k_prv<d_prv)and(k_cur>d_cur)); death=bool((k_prv>d_prv)and(k_cur<d_cur))
    return dict(k=_fp(k_cur),d=_fp(d_cur),kd_golden=golden,kd_death=death,
                kd_oversold=bool(k_cur<20),kd_overbought=bool(k_cur>80),
                kd_k_above_d=bool(k_cur>d_cur),
                kd_signal=('黃金交叉' if golden else '死亡交叉' if death else
                            '超賣區' if k_cur<20 else '多頭排列' if k_cur>d_cur else '空頭排列'),
                _kd_K=K,_kd_D=D)


def calc_macd(close,fast=12,slow=26,sig=9):
    ema_f=close.ewm(span=fast,adjust=False).mean(); ema_s=close.ewm(span=slow,adjust=False).mean()
    ml=ema_f-ema_s; sl=ml.ewm(span=sig,adjust=False).mean(); hs=ml-sl
    m_cur,m_prv=ml.iloc[-1],ml.iloc[-2]; s_cur,s_prv=sl.iloc[-1],sl.iloc[-2]
    h_cur,h_prv=hs.iloc[-1],hs.iloc[-2]
    golden=bool((m_prv<s_prv)and(m_cur>s_cur)); death=bool((m_prv>s_prv)and(m_cur<s_cur))
    return dict(macd=_fp(m_cur),macd_sig_val=_fp(s_cur),macd_hist=_fp(h_cur),
                macd_golden=golden,macd_death=death,macd_above_zero=bool(m_cur>0),
                macd_hist_expand=bool(h_cur>0 and h_cur>h_prv),
                macd_signal_str=('黃金交叉' if golden else '死亡交叉' if death else
                                  '零軸上方' if m_cur>0 else '零軸下方'),
                _macd_line=ml,_macd_sig=sl,_macd_hist=hs)


def calc_bband(close,n=20,k=2):
    mid=close.rolling(n).mean(); std=close.rolling(n).std()
    upper=mid+k*std; lower=mid-k*std
    c=close.iloc[-1]; u,m,l=upper.iloc[-1],mid.iloc[-1],lower.iloc[-1]
    pos=_fp((c-l)/(u-l)) if (u-l)>0 else 0.5
    return dict(bb_upper=_fp(u),bb_mid=_fp(m),bb_lower=_fp(l),bb_pos=pos,
                bb_near_upper=bool(c>=u*0.99),bb_near_lower=bool(c<=l*1.01),
                bb_above_mid=bool(c>m),
                bb_signal=('突破上軌' if c>=u*0.99 else '接近下軌' if c<=l*1.01 else
                            '中軌以上' if c>m else '中軌以下'),
                _bb_upper=upper,_bb_mid=mid,_bb_lower=lower)


def calc_all_ta(hist):
    try:
        if hist is None or len(hist)<30: return {}
        skip={'_kd_K','_kd_D','_macd_line','_macd_sig','_macd_hist','_bb_upper','_bb_mid','_bb_lower'}
        out={}
        for fn,args in [(calc_kd,(hist['High'],hist['Low'],hist['Close'])),
                        (calc_macd,(hist['Close'],)),(calc_bband,(hist['Close'],))]:
            for k,v in fn(*args).items():
                if k not in skip: out[k]=_to_python(v)
        return out
    except Exception as e:
        print(f'  [TA] {e}'); return {}


# ══════════════════════════════════════════════════════
# 均量條件
# ══════════════════════════════════════════════════════
def analyze_volume_condition(vol_series, min_days=2, max_days=5):
    if vol_series is None or len(vol_series)<25: return None
    ma5=vol_series.rolling(5).mean(); ma20=vol_series.rolling(20).mean()
    df=pd.DataFrame({'ma5':ma5,'ma20':ma20}).dropna()
    if df.empty: return None
    consecutive=0
    for v in reversed((df['ma5']>df['ma20']).values):
        if v: consecutive+=1
        else: break
    if not (min_days<=consecutive<=max_days): return None
    ma5_v=int(round(df['ma5'].iloc[-1])); ma20_v=int(round(df['ma20'].iloc[-1]))
    return {'consecutive_days':consecutive,'ma5_volume':ma5_v,'ma20_volume':ma20_v,
            'ratio':float(round(ma5_v/ma20_v,3)) if ma20_v else 0.0}


# ══════════════════════════════════════════════════════
# 背景任務
# ══════════════════════════════════════════════════════
def run_analysis_task(task_id):
    task=tasks[task_id]; task['status']='fetching'; t0=time.time()
    try:
        mode=task.get('mode','both')
        labels={'twse':'上市漲幅前100','tpex':'上櫃漲幅前100','combined':'合併漲幅前100',
                'both':'上市+上櫃各漲幅前100',
                'all_twse':f'全上市({len(_FALLBACK_TWSE)}支)','all_tpex':f'全上櫃({len(_FALLBACK_TPEX)}支)',
                'all_both':f'全市場({len(_FALLBACK_TWSE)+len(_FALLBACK_TPEX)}支)'}
        task['msg']=f"取得【{labels.get(mode,mode)}】排行榜..."
        stocks=get_ranking_stocks(top_n=100,mode=mode)

        data_source_log[task_id]={'mode':mode,'raw_count':len(stocks),
            'min_chg':round(min((s['change_pct'] for s in stocks),default=0),2),
            'max_chg':round(max((s['change_pct'] for s in stocks),default=0),2),
            'timestamp':datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

        if not stocks:
            task.update({'status':'error','msg':'無法取得排行資料（所有 API 離線）'}); return

        total=len(stocks)
        task.update({'total':total,'status':'analyzing','msg':f'取得 {total} 支，計算均量＋技術指標...'})

        qualified=[]
        for i,stock in enumerate(stocks):
            task.update({'current':i+1,'current_code':stock['code'],'current_name':stock['name']})
            hist,actual_sym=get_ohlcv(stock['symbol'],days=90)
            if hist is None:
                task.setdefault('fetch_failed',[]).append(stock['code'])
            vol_ok=analyze_volume_condition(hist['Volume'] if hist is not None else None)
            if vol_ok:
                ta=calc_all_ta(hist)
                real_chg=stock.get('change_pct',0.0)
                if real_chg==0.0 and hist is not None and len(hist)>=2:
                    try:
                        c_now=float(hist['Close'].iloc[-1]); c_prev=float(hist['Close'].iloc[-2])
                        if c_prev>0: real_chg=round((c_now-c_prev)/c_prev*100,2)
                    except Exception: pass
                entry={**stock,'symbol':actual_sym,'change_pct':real_chg,**vol_ok,**ta}
                qualified.append(entry); task['total_found']=len(qualified)
            time.sleep(0.1)

        qualified.sort(key=lambda x:(-x['consecutive_days'],-x['ratio']))
        task.update({'status':'done','stocks':qualified,'total_found':len(qualified),
                     'scanned':total,'elapsed':round(time.time()-t0,1),
                     'timestamp':datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                     'msg':f'完成！找到 {len(qualified)} 支符合均量條件個股'})
    except Exception as e:
        import traceback; print(traceback.format_exc())
        task.update({'status':'error','msg':f'分析錯誤：{e}'})


# ══════════════════════════════════════════════════════
# Flask Routes
# ══════════════════════════════════════════════════════
@app.route('/')
def index():
    static_dir=os.path.join(os.path.dirname(os.path.abspath(__file__)),'static')
    return send_from_directory(static_dir,'index.html')

@app.route('/api/debug')
def debug_info():
    recent={}
    for tid,log in list(data_source_log.items())[-5:]:
        task=tasks.get(tid,{}); failed=task.get('fetch_failed',[])
        recent[tid]={**log,'task_status':task.get('status'),
                     'stocks_found':task.get('total_found',0),
                     'fetch_failed_count':len(failed),'fetch_failed_codes':failed[:20]}
    return jsonify({'recent_scans':recent,
                    'hints':{'fetch_failed':'yfinance取不到歷史的代碼→均量無法計算→可能漏選',
                             'min_chg_low':'min_chg<1%代表Yahoo Screener失敗，使用yfinance備援'}})

@app.route('/api/analyze/start',methods=['POST','GET'])
def analyze_start():
    body=request.get_json(silent=True) or {}
    mode=body.get('mode','both')
    if mode not in ('twse','tpex','combined','both','all_twse','all_tpex','all_both'):
        mode='both'
    task_id=str(uuid.uuid4())[:8]
    tasks[task_id]={'status':'pending','msg':'準備開始...','mode':mode,
                    'current':0,'total':0,'current_code':'','current_name':'',
                    'stocks':[],'total_found':0,'scanned':0,'elapsed':0,'timestamp':''}
    threading.Thread(target=run_analysis_task,args=(task_id,),daemon=True).start()
    return jsonify({'task_id':task_id})

@app.route('/api/analyze/status/<task_id>')
def analyze_status(task_id):
    task=tasks.get(task_id)
    if not task: return jsonify({'error':'找不到任務'}),404
    return jsonify(task)

@app.route('/api/stock/<code>')
def stock_detail(code):
    hist,sym=get_ohlcv(f'{code}.TW',90)
    if hist is None: hist,sym=get_ohlcv(f'{code}.TWO',90)
    if hist is None: return jsonify({'error':f'無法取得 {code} 資料'}),404
    close,high,low,vol=hist['Close'],hist['High'],hist['Low'],hist['Volume']
    dates=[d.strftime('%Y-%m-%d') for d in hist.index]
    def s(series): return [_fp(v) for v in series]
    def iv(series): return [int(v) if pd.notna(v) else None for v in series]
    kd_d=calc_kd(high,low,close); macd_d=calc_macd(close); bb_d=calc_bband(close)
    return jsonify({
        'code':code,'symbol':sym,'dates':dates,
        'open':s(hist['Open']),'high':s(high),'low':s(low),'close':s(close),
        'volume':iv(vol),'ma5v':iv(vol.rolling(5).mean()),'ma20v':iv(vol.rolling(20).mean()),
        'ma5p':s(close.rolling(5).mean()),'ma20p':s(close.rolling(20).mean()),
        'macd':s(macd_d['_macd_line']),'macd_signal':s(macd_d['_macd_sig']),'macd_hist':s(macd_d['_macd_hist']),
        'kd_k':s(kd_d['_kd_K']),'kd_d':s(kd_d['_kd_D']),
        'bb_upper':s(bb_d['_bb_upper']),'bb_mid':s(bb_d['_bb_mid']),'bb_lower':s(bb_d['_bb_lower']),
    })


# ══════════════════════════════════════════════════════
# 匯出 PDF
# ══════════════════════════════════════════════════════
@app.route('/api/export/pdf',methods=['POST'])
def export_pdf():
    from reportlab.platypus import SimpleDocTemplate,Table,TableStyle,Paragraph,Spacer,HRFlowable
    from reportlab.lib.pagesizes import A4,landscape
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    body=request.get_json(force=True)
    stocks=body.get('stocks',[]); ts=body.get('timestamp',datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    scanned=body.get('scanned',0); mode_label=body.get('mode_label','')
    buf=io.BytesIO()
    doc=SimpleDocTemplate(buf,pagesize=landscape(A4),leftMargin=12*mm,rightMargin=12*mm,topMargin=14*mm,bottomMargin=14*mm)
    FONT='Helvetica'
    try:
        import glob
        cands=glob.glob('/usr/share/fonts/truetype/noto/*CJK*Regular*.ttf')+glob.glob('/usr/share/fonts/truetype/wqy/*.ttf')
        if cands: pdfmetrics.registerFont(TTFont('CJK',cands[0])); FONT='CJK'
    except Exception: pass
    def PS(name,**kw): return ParagraphStyle(name,fontName=FONT,**kw)
    def fv(n):
        if not n: return '-'
        if n>=1e8: return f'{n/1e8:.1f}億'
        if n>=1e4: return f'{n/1e3:.0f}K'
        return str(n)
    story=[Paragraph('台股均量追蹤系統 掃描報告',PS('T',fontSize=15,textColor=colors.HexColor('#00d4aa'),spaceAfter=4)),
           Paragraph(f'掃描時間：{ts}　範圍：{mode_label}　已掃描：{scanned}支　符合：{len(stocks)}支',PS('S',fontSize=8,textColor=colors.HexColor('#94a3b8'),spaceAfter=6)),
           HRFlowable(width='100%',thickness=0.8,color=colors.HexColor('#1e2d47'),spaceAfter=8)]
    if stocks:
        hdrs=['代碼','名稱','市場','漲幅','連續','5MA量','比率','K','D','KD訊號','MACD','柱狀','MACD狀態','布林位','布林訊號']
        rows=[hdrs]
        for s in stocks:
            chg=f"+{s.get('change_pct',0):.2f}%" if s.get('change_pct',0)>=0 else f"{s.get('change_pct',0):.2f}%"
            rows.append([s.get('code',''),s.get('name',''),s.get('market',''),chg,str(s.get('consecutive_days','')),
                         fv(s.get('ma5_volume')),f"{s.get('ratio',0):.2f}x",
                         f"{s.get('k',0):.1f}" if s.get('k') is not None else '-',
                         f"{s.get('d',0):.1f}" if s.get('d') is not None else '-',
                         s.get('kd_signal','-'),
                         f"{s.get('macd',0):.3f}" if s.get('macd') is not None else '-',
                         f"{s.get('macd_hist',0):.3f}" if s.get('macd_hist') is not None else '-',
                         s.get('macd_signal_str','-'),
                         f"{s.get('bb_pos',0):.0%}" if s.get('bb_pos') is not None else '-',
                         s.get('bb_signal','-')])
        cw=[c*mm for c in [14,28,12,14,13,16,14,11,11,18,16,16,18,14,18]]
        t=Table(rows,colWidths=cw,repeatRows=1)
        t.setStyle(TableStyle([('FONTNAME',(0,0),(-1,-1),FONT),('FONTSIZE',(0,0),(-1,-1),7),
            ('ALIGN',(0,0),(-1,-1),'CENTER'),('VALIGN',(0,0),(-1,-1),'MIDDLE'),
            ('BACKGROUND',(0,0),(-1,0),colors.HexColor('#0f3460')),('TEXTCOLOR',(0,0),(-1,0),colors.HexColor('#00d4aa')),
            ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.HexColor('#111827'),colors.HexColor('#0f172a')]),
            ('TEXTCOLOR',(0,1),(-1,-1),colors.HexColor('#cbd5e1')),('GRID',(0,0),(-1,-1),0.25,colors.HexColor('#1e2d47')),
            ('TOPPADDING',(0,0),(-1,-1),4),('BOTTOMPADDING',(0,0),(-1,-1),4),
            ('TEXTCOLOR',(3,1),(3,-1),colors.HexColor('#00d98b'))]))
        story.append(t)
    story+=[Spacer(1,6*mm),HRFlowable(width='100%',thickness=0.4,color=colors.HexColor('#1e2d47'),spaceAfter=3),
            Paragraph('資料來源：Yahoo Finance/TWSE/TPEX　本報告僅供參考，不構成投資建議',PS('F',fontSize=7,textColor=colors.HexColor('#64748b')))]
    doc.build(story); buf.seek(0)
    return send_file(buf,mimetype='application/pdf',as_attachment=True,download_name=f"stock_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf")


# ══════════════════════════════════════════════════════
# 匯出 DOCX
# ══════════════════════════════════════════════════════
@app.route('/api/export/docx',methods=['POST'])
def export_docx():
    from docx import Document
    from docx.shared import Pt,RGBColor,Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    body=request.get_json(force=True)
    stocks=body.get('stocks',[]); ts=body.get('timestamp',datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    scanned=body.get('scanned',0); mode_label=body.get('mode_label','')
    doc=Document(); sec=doc.sections[0]
    sec.page_width=Cm(42); sec.page_height=Cm(29.7); sec.left_margin=sec.right_margin=Cm(1.2); sec.top_margin=sec.bottom_margin=Cm(1.5)
    def set_bg(cell,hc):
        pr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
        shd.set(qn('w:fill'),hc); shd.set(qn('w:val'),'clear'); pr.append(shd)
    def set_border(cell):
        pr=cell._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
        for side in ['top','left','bottom','right']:
            el=OxmlElement(f'w:{side}'); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4'); el.set(qn('w:color'),'1E2D47'); b.append(el)
        pr.append(b)
    h=doc.add_heading('台股均量追蹤系統 — 掃描報告',0); h.alignment=WD_ALIGN_PARAGRAPH.LEFT
    r=h.runs[0]; r.font.color.rgb=RGBColor(0,0xD4,0xAA); r.font.size=Pt(16)
    p=doc.add_paragraph()
    for t in [f'掃描時間：{ts}',f'　範圍：{mode_label}',f'　已掃描：{scanned}支',f'　符合：{len(stocks)}支']:
        rr=p.add_run(t); rr.font.size=Pt(9); rr.font.color.rgb=RGBColor(0x94,0xA3,0xB8)
    doc.add_paragraph()
    def fv(n):
        if not n: return '-'
        if n>=1e8: return f'{n/1e8:.1f}億'
        if n>=1e4: return f'{n/1e3:.0f}K'
        return str(n)
    hdrs=['代碼','名稱','市場','漲幅%','連續','5MA量','比率','K','D','KD訊號','MACD','柱狀','MACD狀態','布林位','布林訊號']
    widths=[1.6,3.2,1.2,1.6,1.4,2.2,1.6,1.3,1.3,2.2,2.0,2.0,2.2,1.8,2.2]
    tbl=doc.add_table(rows=1,cols=len(hdrs)); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.style='Table Grid'
    for i,w in enumerate(widths):
        for cell in tbl.columns[i].cells: cell.width=Cm(w)
    for cell,txt in zip(tbl.rows[0].cells,hdrs):
        set_bg(cell,'0F3460'); set_border(cell); cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(txt); run.bold=True; run.font.size=Pt(8); run.font.color.rgb=RGBColor(0,0xD4,0xAA)
    for idx,s in enumerate(stocks):
        row=tbl.add_row(); row.height=Cm(0.72); bg='111827' if idx%2==0 else '0F172A'
        chg=f"+{s.get('change_pct',0):.2f}%" if s.get('change_pct',0)>=0 else f"{s.get('change_pct',0):.2f}%"
        vals=[s.get('code',''),s.get('name',''),s.get('market',''),chg,str(s.get('consecutive_days','')),
              fv(s.get('ma5_volume')),f"{s.get('ratio',0):.2f}x",
              f"{s.get('k',0):.1f}" if s.get('k') is not None else '-',
              f"{s.get('d',0):.1f}" if s.get('d') is not None else '-',
              s.get('kd_signal','-'),
              f"{s.get('macd',0):.3f}" if s.get('macd') is not None else '-',
              f"{s.get('macd_hist',0):.3f}" if s.get('macd_hist') is not None else '-',
              s.get('macd_signal_str','-'),
              f"{s.get('bb_pos',0):.0%}" if s.get('bb_pos') is not None else '-',
              s.get('bb_signal','-')]
        vcols=[RGBColor(0xFF,0xFF,0xFF),RGBColor(0xCB,0xD5,0xE1),RGBColor(0xFF,0xB5,0x47),
               (RGBColor(0,0xD9,0x8B) if s.get('change_pct',0)>=0 else RGBColor(0xFF,0x4D,0x6D)),
               RGBColor(0,0xD4,0xAA),RGBColor(0,0xD4,0xAA),RGBColor(0,0x84,0xFF),
               RGBColor(0xFF,0xD7,0),RGBColor(0xFF,0xD7,0),RGBColor(0xCB,0xD5,0xE1),
               RGBColor(0,0x84,0xFF),RGBColor(0x94,0xA3,0xB8),RGBColor(0xCB,0xD5,0xE1),
               RGBColor(0x7C,0x3A,0xED),RGBColor(0xCB,0xD5,0xE1)]
        for cell,val,vc in zip(row.cells,vals,vcols):
            set_bg(cell,bg); set_border(cell); cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            run=p.add_run(val); run.font.size=Pt(7.5); run.font.color.rgb=vc
    doc.add_paragraph()
    fp=doc.add_paragraph('資料來源：Yahoo Finance/TWSE/TPEX　本報告僅供參考，不構成投資建議')
    fp.runs[0].font.size=Pt(7.5); fp.runs[0].font.color.rgb=RGBColor(0x64,0x74,0x8B)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return send_file(buf,mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True,download_name=f"stock_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")


if __name__ == '__main__':
    os.makedirs('static', exist_ok=True)
    print('='*55)
    print('🚀 台股均量追蹤系統 v5')
    print(f'📡 上市 {len(_FALLBACK_TWSE)} 支 + 上櫃 {len(_FALLBACK_TPEX)} 支')
    print('🌐 http://localhost:5001')
    print('='*55)
    app.run(host='0.0.0.0', port=5001, debug=False, threaded=True)
